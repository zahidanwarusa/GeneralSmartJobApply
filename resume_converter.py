"""
Resume Converter - Convert any resume format to JSON

This module converts resumes in various formats (PDF, DOCX, TXT) to the structured
JSON format used by SmartApplyPro, with AI-powered extraction and optimization.

Usage:
    python resume_converter.py --input resume.pdf --output resume.json
    python resume_converter.py --input resume.docx --optimize --job-file data/jobs/job.json
    python resume_converter.py --input resume.txt --output custom_resume.json --name "John Doe"
"""

import json
import logging
import os
import re
import argparse
from pathlib import Path
from typing import Dict, List, Optional
import google.generativeai as genai
from config import GEMINI_API_KEYS, DATA_DIR, DEFAULT_RESUME
from gemini_service import GeminiService

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ResumeConverter:
    """Convert resumes from various formats to structured JSON"""

    def __init__(self):
        """Initialize the resume converter with Gemini service"""
        self.gemini_service = GeminiService()
        self.logger = logging.getLogger(__name__)

        # Ensure output directory exists
        self.output_dir = DATA_DIR / 'resumes'
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path.suffix.lower()

        try:
            if file_extension == '.txt':
                return self._extract_from_txt(file_path)
            elif file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file using PyPDF2"""
        try:
            import PyPDF2

            text = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())

            return '\n'.join(text)
        except ImportError:
            self.logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise ImportError("PyPDF2 is required for PDF extraction. Install with: pip install PyPDF2")
        except Exception as e:
            self.logger.error(f"Error extracting PDF: {str(e)}")
            raise

    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file using python-docx"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = []

            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)

            return '\n'.join(text)
        except ImportError:
            self.logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
        except Exception as e:
            self.logger.error(f"Error extracting DOCX: {str(e)}")
            raise

    def convert_to_json(self, resume_text: str, candidate_name: Optional[str] = None) -> Dict:
        """Convert resume text to structured JSON using Gemini AI"""

        # Load the default template to show structure
        with open(DEFAULT_RESUME, 'r', encoding='utf-8') as f:
            template = json.load(f)

        # Create prompt for Gemini
        prompt = f"""
        Convert the following resume into a structured JSON format matching this exact structure:

        TEMPLATE STRUCTURE:
        {json.dumps(template, indent=2)}

        RESUME TEXT:
        {resume_text}

        INSTRUCTIONS:
        1. Extract all information from the resume and organize it according to the template structure
        2. The JSON must have these exact sections:
           - header: {{name, email, phone, citizenship (if mentioned), linkedin}}
           - professional_summary: {{title_experience, track_record, expertise, core_value}}
           - core_competencies: {{programming_and_automation, testing_frameworks, cloud_and_devops, api_and_performance, quality_tools, databases, domain_expertise, leadership}}
           - professional_experience: [{{company, location, position, duration, summary, key_achievements[], detailed_achievements[], environment}}]
           - education: [{{degree, major, university, year}}]

        3. For professional_summary:
           - title_experience: Start with role/title and years of experience
           - track_record: Quantifiable achievements and impact
           - expertise: Key technical areas and domains
           - core_value: Value proposition or unique strength

        4. For core_competencies:
           - Organize skills into the 8 categories shown in template
           - Each category should be an array of skills
           - Extract all technical skills, tools, and technologies mentioned

        5. For professional_experience:
           - summary: Brief overview of role and responsibilities
           - key_achievements: 3-5 bullet points of major accomplishments (use **bold** for important terms)
           - detailed_achievements: More detailed accomplishments and contributions
           - environment: Comma-separated list of technologies used (ONLY technical tools/technologies)

        6. For education:
           - Include all degrees with degree type, major, university, and year

        7. Use **bold markers** around important skills and achievements (Example: **Python**, **AWS**, **80% automation**)

        8. If candidate name is provided, use it. Otherwise extract from resume.

        9. If certain information is not in the resume, use reasonable placeholders:
           - Missing email: "email@example.com"
           - Missing phone: "(000) 000-0000"
           - Missing citizenship: "Not Specified"

        10. CRITICAL: Return ONLY valid JSON that can be parsed directly, no markdown, no explanations

        CANDIDATE NAME (if provided): {candidate_name if candidate_name else "Extract from resume"}

        Return the complete JSON object now:
        """

        # Call Gemini API
        self.logger.info("Sending resume to Gemini AI for structured extraction...")
        response = self.gemini_service.make_api_call(
            prompt,
            generation_config=genai.GenerationConfig(
                temperature=0.1,
                top_p=1,
                top_k=1,
                max_output_tokens=8000,
            )
        )

        if not response or not hasattr(response, 'text'):
            raise Exception("Failed to get response from Gemini AI")

        # Extract JSON from response
        resume_json = self._extract_and_validate_json(response.text, template)

        # Override name if provided
        if candidate_name:
            resume_json['header']['name'] = candidate_name

        return resume_json

    def _extract_and_validate_json(self, response_text: str, template: Dict) -> Dict:
        """Extract and validate JSON from Gemini response"""

        # Try to extract JSON
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            json_text = json_match.group(0)

            # Clean up
            json_text = re.sub(r'```(?:json)?', '', json_text)
            json_text = re.sub(r'```', '', json_text)
            json_text = re.sub(r',\s*}', '}', json_text)
            json_text = re.sub(r',\s*]', ']', json_text)

            try:
                resume_json = json.loads(json_text)
            except json.JSONDecodeError as e:
                self.logger.error(f"Failed to parse JSON: {str(e)}")
                # Save problematic response
                with open(self.output_dir / 'failed_extraction.txt', 'w') as f:
                    f.write(response_text)
                raise Exception("Failed to parse JSON from AI response")
        else:
            raise Exception("No JSON found in AI response")

        # Validate structure
        required_sections = ['header', 'professional_summary', 'core_competencies',
                           'professional_experience', 'education']

        missing_sections = [s for s in required_sections if s not in resume_json]
        if missing_sections:
            self.logger.warning(f"Missing sections: {missing_sections}")

            # Add missing sections from template
            for section in missing_sections:
                resume_json[section] = template[section]

        # Validate header fields
        required_header_fields = ['name', 'email', 'phone']
        for field in required_header_fields:
            if field not in resume_json['header']:
                resume_json['header'][field] = f"Missing {field}"

        # Validate professional_summary fields
        required_summary_fields = ['title_experience', 'track_record', 'expertise', 'core_value']
        for field in required_summary_fields:
            if field not in resume_json['professional_summary']:
                resume_json['professional_summary'][field] = "Not specified"

        # Ensure core_competencies has all categories
        required_competency_categories = [
            'programming_and_automation', 'testing_frameworks', 'cloud_and_devops',
            'api_and_performance', 'quality_tools', 'databases',
            'domain_expertise', 'leadership'
        ]
        for category in required_competency_categories:
            if category not in resume_json['core_competencies']:
                resume_json['core_competencies'][category] = []

        # Ensure professional_experience is a list
        if not isinstance(resume_json['professional_experience'], list):
            resume_json['professional_experience'] = [resume_json['professional_experience']]

        # Ensure education is a list
        if not isinstance(resume_json['education'], list):
            resume_json['education'] = [resume_json['education']]

        return resume_json

    def optimize_with_job_description(self, resume_json: Dict, job_file: str) -> Dict:
        """Optimize resume JSON based on job description"""

        # Load job description
        job_file_path = Path(job_file)
        if not job_file_path.exists():
            raise FileNotFoundError(f"Job file not found: {job_file}")

        with open(job_file_path, 'r', encoding='utf-8') as f:
            job_details = json.load(f)

        self.logger.info(f"Optimizing resume for job: {job_details.get('title', 'Unknown')}")

        # Optimize each section
        optimized = resume_json.copy()

        # Optimize professional summary
        self.logger.info("Optimizing professional summary...")
        optimized['professional_summary'] = self.gemini_service.optimize_resume_section(
            'professional_summary',
            resume_json['professional_summary'],
            job_details
        )

        # Optimize core competencies
        self.logger.info("Optimizing core competencies...")
        optimized['core_competencies'] = self.gemini_service.optimize_resume_section(
            'core_competencies',
            resume_json['core_competencies'],
            job_details
        )

        # Optimize professional experience
        self.logger.info("Optimizing professional experience...")
        optimized['professional_experience'] = self.gemini_service.optimize_resume_section(
            'professional_experience',
            resume_json['professional_experience'],
            job_details
        )

        return optimized

    def save_json(self, resume_json: Dict, output_file: str) -> str:
        """Save resume JSON to file"""

        output_path = Path(output_file)

        # Ensure directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save with proper formatting
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(resume_json, f, indent=2, ensure_ascii=False)

        self.logger.info(f"Resume saved to: {output_path}")
        return str(output_path)

    def convert_file(self, input_file: str, output_file: Optional[str] = None,
                    candidate_name: Optional[str] = None,
                    optimize: bool = False,
                    job_file: Optional[str] = None) -> str:
        """
        Complete workflow: Extract text, convert to JSON, optionally optimize, and save

        Args:
            input_file: Path to resume file (PDF, DOCX, or TXT)
            output_file: Path to output JSON file (optional)
            candidate_name: Name to use in resume (optional)
            optimize: Whether to optimize with job description
            job_file: Path to job description JSON (required if optimize=True)

        Returns:
            Path to saved JSON file
        """

        # Extract text
        self.logger.info(f"Extracting text from: {input_file}")
        resume_text = self.extract_text_from_file(input_file)

        # Convert to JSON
        self.logger.info("Converting to JSON format...")
        resume_json = self.convert_to_json(resume_text, candidate_name)

        # Optimize if requested
        if optimize:
            if not job_file:
                raise ValueError("job_file is required when optimize=True")
            resume_json = self.optimize_with_job_description(resume_json, job_file)

        # Determine output file
        if not output_file:
            input_path = Path(input_file)
            name = resume_json['header'].get('name', 'resume').replace(' ', '_')
            output_file = self.output_dir / f"{name}_resume.json"

        # Save JSON
        return self.save_json(resume_json, output_file)


def main():
    """Command-line interface"""
    parser = argparse.ArgumentParser(
        description='Convert resume files to JSON format with AI-powered extraction',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Basic conversion
  python resume_converter.py --input resume.pdf

  # Specify output file
  python resume_converter.py --input resume.docx --output my_resume.json

  # With candidate name
  python resume_converter.py --input resume.txt --name "John Doe"

  # With job description optimization
  python resume_converter.py --input resume.pdf --optimize --job-file data/jobs/sdet_job.json

  # Complete workflow
  python resume_converter.py --input resume.pdf --output custom.json --name "Jane Smith" --optimize --job-file data/jobs/job.json
        """
    )

    parser.add_argument('--input', '-i', required=True, help='Input resume file (PDF, DOCX, or TXT)')
    parser.add_argument('--output', '-o', help='Output JSON file path (optional)')
    parser.add_argument('--name', '-n', help='Candidate name to use in resume (optional)')
    parser.add_argument('--optimize', action='store_true', help='Optimize resume with job description')
    parser.add_argument('--job-file', '-j', help='Job description JSON file for optimization')

    args = parser.parse_args()

    # Validate arguments
    if args.optimize and not args.job_file:
        parser.error("--job-file is required when using --optimize")

    # Run conversion
    try:
        converter = ResumeConverter()
        output_path = converter.convert_file(
            input_file=args.input,
            output_file=args.output,
            candidate_name=args.name,
            optimize=args.optimize,
            job_file=args.job_file
        )

        print("\n" + "="*60)
        print("✓ Resume conversion completed successfully!")
        print("="*60)
        print(f"Output file: {output_path}")
        print("\nYou can now use this JSON file with SmartApplyPro:")
        print(f"  python main.py --mode resume --job-file data/jobs/job.json")
        print("="*60 + "\n")

    except Exception as e:
        logger.error(f"Conversion failed: {str(e)}")
        print(f"\n❌ Error: {str(e)}\n")
        return 1

    return 0


if __name__ == '__main__':
    exit(main())
